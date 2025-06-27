"""
Utilidades para exportación de informes en PDF y Word
Autor: Valentina Bailon Cano
"""

from fpdf import FPDF
from docx import Document

def exportar_pdf(texto: str, nombre_archivo: str):
    """
    Genera un informe en formato PDF desde texto plano.
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for linea in texto.splitlines():
        pdf.multi_cell(0, 10, linea)
    ruta = f"outputs/{nombre_archivo}"
    pdf.output(ruta)
    return ruta

def exportar_docx(texto: str, nombre_archivo: str):
    """
    Genera un informe en formato Word (.docx).
    """
    doc = Document()
    doc.add_heading("Informe Institucional EURO STOXX 50", 0)
    for linea in texto.splitlines():
        doc.add_paragraph(linea)
    ruta = f"outputs/{nombre_archivo}"
    doc.save(ruta)
    return ruta
